"""
Parser de CVs : PDF (natif + scanné) et DOCX → texte propre.

Stratégie par défaut :
1. PDF        → pypdf
2. Fallback   → pytesseract OCR
3. Escalade   → OCR Vision OpenAI-compatible (ex: Qwen local)
4. DOCX       → python-docx extraction paragraphes + tableaux
5. Nettoyage  → suppression artefacts, normalisation espaces/sauts de ligne
"""

import base64
import logging
import os
import re
import tempfile
from pathlib import Path

from pypdf import PdfReader
from docx import Document
from openai import OpenAI
import numpy as np

logger = logging.getLogger(__name__)

# Seuils de qualité du texte extrait
MIN_CHARS_PER_PAGE = int(os.getenv("MIN_CHARS_PER_PAGE", "30"))
OCR_MIN_CHARS = int(os.getenv("OCR_MIN_CHARS", "200"))
OCR_MIN_WORDS = int(os.getenv("OCR_MIN_WORDS", "40"))
OCR_MIN_ALPHA_RATIO = float(os.getenv("OCR_MIN_ALPHA_RATIO", "0.60"))
OCR_MIN_WORDS_PER_PAGE = int(os.getenv("OCR_MIN_WORDS_PER_PAGE", "20"))
OCR_MIN_ANCHOR_HITS = int(os.getenv("OCR_MIN_ANCHOR_HITS", "1"))
OCR_ENABLE_LLM_FALLBACK = os.getenv("OCR_ENABLE_LLM_FALLBACK", "true").lower() == "true"
PDF_PARSE_MODE = os.getenv("PDF_PARSE_MODE", "smart").lower()
LLM_VISION_MODEL = os.getenv("LLM_VISION_MODEL", "gpt-4o")


def parse_file(file_path: str | Path) -> dict:
    """Point d'entrée principal. Retourne un dict avec le texte et des métadonnées."""
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")

    # CONFIGURATION
    parse_mode = os.getenv("PDF_PARSE_MODE", "smart").lower()
    suffix = file_path.suffix.lower()
    text = ""
    method = ""
    disable_docling = os.getenv("DISABLE_DOCLING", "0") == "1"

    # --- STRATÉGIE POUR LES PDF ---
    if suffix == ".pdf":
        # 0. Mode colonnes explicite
        if parse_mode in ["pymupdf-columns", "columns"]:
            text, method = _parse_pdf_pymupdf_columns(file_path)
            quality = _compute_text_quality(text)
            if not _is_quality_sufficient(quality):
                text = ""

        # 1. Mode auto/smart : extraction colonnes si layout 2 colonnes
        if not text and parse_mode in ["auto", "smart"]:
            if _is_two_columns_pdf(file_path):
                logger.info("Colonnes détectées : extraction PyMuPDF (colonnes).")
                text, method = _parse_pdf_pymupdf_columns(file_path)
                quality = _compute_text_quality(text)
                if not _is_quality_sufficient(quality):
                    text = ""

        # 2. Mode docling explicite seulement
        if not text and parse_mode == "docling" and not disable_docling:
            try:
                logger.info("Tentative de parsing via Docling : %s", file_path.name)
                text = _parse_via_docling(file_path)
                quality = _compute_text_quality(text)
                if _is_quality_sufficient(quality):
                    method = "docling-markdown"
                else:
                    logger.warning("Docling a produit un texte de faible qualité. Passage au pipeline OCR standard...")
                    text = ""
            except Exception as e:
                logger.warning("Docling a échoué : %s. Passage au pipeline OCR standard...", e)

        # 3. Pipeline OCR standard : pypdf -> tesseract -> vision local/distant
        if not text:
            text, method = _parse_pdf(file_path)

    # --- STRATÉGIE POUR LES DOCX ---
    elif suffix in (".docx", ".doc"):
        text, method = _parse_docx(file_path), "python-docx"
    
    else:
        raise ValueError(f"Format non supporté : {suffix}. Formats acceptés : .pdf, .docx")

    # Nettoyage et Qualité
    text = _clean_text(text, is_markdown=(method == "docling-markdown"))
    final_quality = _compute_text_quality(text)

    result = {
        "source": file_path.name,
        "format": suffix.lstrip("."),
        "method": method,
        "char_count": len(text),
        "quality": final_quality,
        "text": text,
    }

    logger.info(
        "Parsing terminé : %s | méthode=%s | %d caractères | alpha_ratio=%.2f",
        file_path.name,
        method,
        len(text),
        final_quality["alpha_ratio"],
    )

    return result


def _parse_pdf_native(file_path: Path) -> tuple[str, str]:
    """Extraction brute via pypdf sans OCR."""
    try:
        reader = PdfReader(str(file_path))
        pages_text = [page.extract_text() or "" for page in reader.pages]
        full_text = "\n".join(pages_text)
        return full_text, "pypdf"
    except Exception as e:
        logger.warning("Erreur pypdf : %s", e)
        return "", ""


def _kmeans_1d(values: list[float], iters: int = 10) -> list[float]:
    if not values:
        return []
    v = sorted(values)
    c1 = v[len(v) // 4]
    c2 = v[(3 * len(v)) // 4]
    for _ in range(iters):
        g1, g2 = [], []
        for x in v:
            if abs(x - c1) <= abs(x - c2):
                g1.append(x)
            else:
                g2.append(x)
        if g1:
            c1 = sum(g1) / len(g1)
        if g2:
            c2 = sum(g2) / len(g2)
    return [c1, c2]


def _is_two_columns_pdf(file_path: Path) -> bool:
    """Heuristique simple pour détecter 2 colonnes via positions X."""
    try:
        import fitz
    except Exception:
        return False

    doc = fitz.open(str(file_path))
    xs: list[float] = []
    for page in doc:
        blocks = page.get_text("blocks")
        for b in blocks:
            if b[4].strip():
                xs.append(float(b[0]))
        if len(xs) > 200:
            break
    doc.close()

    if len(xs) < 10:
        return False

    centers = _kmeans_1d(xs)
    if len(centers) < 2:
        return False
    return abs(centers[0] - centers[1]) > 80


def _parse_pdf_pymupdf_columns(file_path: Path) -> tuple[str, str]:
    """Extraction PyMuPDF avec regroupement par colonnes (gauche puis droite)."""
    try:
        import fitz
    except Exception as e:
        logger.warning("PyMuPDF requis pour l'extraction colonnes : %s", e)
        return "", ""

    doc = fitz.open(str(file_path))
    out_lines = []
    for page in doc:
        blocks = [b for b in page.get_text("blocks") if b[4].strip()]
        if not blocks:
            continue
        xs = [b[0] for b in blocks]
        centers = _kmeans_1d(xs)
        if len(centers) < 2:
            ordered = sorted(blocks, key=lambda b: (b[1], b[0]))
        else:
            split = sum(centers) / 2
            left = [b for b in blocks if b[0] <= split]
            right = [b for b in blocks if b[0] > split]
            left.sort(key=lambda b: (b[1], b[0]))
            right.sort(key=lambda b: (b[1], b[0]))
            ordered = left + right
        out_lines.extend([b[4].strip() for b in ordered])
        out_lines.append("")
    doc.close()
    return "\n".join(out_lines).strip(), "pymupdf-columns"


def _parse_via_docling(file_path: Path) -> str:
    """
    Utilise Docling pour extraire le contenu en Markdown structuré.

    Configuration optimisée pour les CVs :
    - OCR activé pour les PDFs scannés (EasyOCR)
    - Résolution 300 DPI pour l'OCR
    - Reconstruction des tableaux (langues, compétences en tableau)
    - Pour les PDFs natifs, Docling lit directement les vecteurs PDF
      (aucune binarisation n'est appliquée ni utile dans ce cas)
    """
    try:
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.base_models import InputFormat

        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = True                  # Active l'OCR pour les scans
        pipeline_options.do_table_structure = True      # Reconstruit les tableaux (langues, compétences)
        # pipeline_options.ocr_options.use_gpu = False  # Supprimé car cause des erreurs selon la version
        pipeline_options.images_scale = 2.0             # Résolution ×2 pour l'OCR interne (~300 DPI)

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        result = converter.convert(str(file_path))
        return result.document.export_to_markdown()

    except ImportError:
        raise ImportError("Docling n'est pas installé. pip install docling")
    except TypeError:
        # Ancienne version de docling sans PdfPipelineOptions — fallback basique
        logger.warning("Version de Docling sans PdfPipelineOptions — configuration basique utilisée.")
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(str(file_path))
        return result.document.export_to_markdown()



# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(file_path: Path) -> tuple[str, str]:
    """Extrait le texte d'un PDF via pypdf -> pytesseract -> OCR Vision."""
    reader = PdfReader(str(file_path))
    pages_text = []

    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)

    full_text = "\n".join(pages_text)
    page_count = max(len(pages_text), 1)
    native_quality = _compute_text_quality(full_text, page_count=page_count)

    if _is_quality_sufficient(native_quality):
        logger.info(
            "Extraction pypdf suffisante (%d car/page, %d mots/page, alpha_ratio=%.2f, anchors=%d)",
            int(native_quality["chars_per_page"]),
            int(native_quality["words_per_page"]),
            native_quality["alpha_ratio"],
            native_quality["anchor_hits"],
        )
        return full_text, "pypdf"

    logger.info(
        "Texte pypdf insuffisant (%d car/page, %d mots/page, alpha_ratio=%.2f). Tentative Tesseract...",
        int(native_quality["chars_per_page"]),
        int(native_quality["words_per_page"]),
        native_quality["alpha_ratio"],
    )
    ocr_text = _ocr_pdf(file_path)
    ocr_quality = _compute_text_quality(ocr_text, page_count=page_count)
    logger.info(
        "Résultat Tesseract : char_count=%d, word_count=%d, alpha_ratio=%.2f, anchors=%d",
        ocr_quality["char_count"],
        ocr_quality["word_count"],
        ocr_quality["alpha_ratio"],
        ocr_quality["anchor_hits"],
    )

    if _is_quality_sufficient(ocr_quality):
        return ocr_text, "pytesseract-ocr"

    best_text, best_method, best_quality = _best_quality_candidate(
        [
            (full_text, "pypdf", native_quality),
            (ocr_text, "pytesseract-ocr", ocr_quality),
        ]
    )

    if PDF_PARSE_MODE in {"api", "api-vision", "vision", "smart", "auto", "qwen", "qwen-local"} and OCR_ENABLE_LLM_FALLBACK:
        try:
            logger.info("Qualité encore insuffisante après Tesseract. Tentative OCR Vision/Qwen...")
            vision_text = _ocr_pdf_via_api(file_path)
            vision_quality = _compute_text_quality(vision_text, page_count=page_count)
            logger.info(
                "Résultat Vision : char_count=%d, word_count=%d, alpha_ratio=%.2f, anchors=%d",
                vision_quality["char_count"],
                vision_quality["word_count"],
                vision_quality["alpha_ratio"],
                vision_quality["anchor_hits"],
            )
            if _is_quality_sufficient(vision_quality):
                return vision_text, "llm-vision-api"
            best_text, best_method, best_quality = _best_quality_candidate(
                [
                    (best_text, best_method, best_quality),
                    (vision_text, "llm-vision-api", vision_quality),
                ]
            )
        except Exception as e:
            logger.warning("Vision API/Qwen indisponible (%s). Conservation du meilleur fallback local.", e)

    logger.warning(
        "Aucune extraction n'atteint le seuil cible; conservation de %s (score composite le plus élevé).",
        best_method,
    )
    return best_text, best_method


def _get_vision_client() -> tuple[OpenAI, str]:
    """Crée un client OpenAI-compatible pour l'OCR Vision."""
    provider = os.getenv("OCR_PROVIDER", os.getenv("LLM_PROVIDER", "copilot")).lower()

    if provider == "gemini":
        api_key = os.getenv("GEMINI_API_KEY", "")
        base_url = "https://generativelanguage.googleapis.com/v1beta/openai/"
        model = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")
    elif provider == "openai":
        api_key = os.getenv("OPENAI_API_KEY", "")
        base_url = os.getenv("OPENAI_BASE_URL", "")
        model = os.getenv("OPENAI_MODEL", LLM_VISION_MODEL)
    elif provider == "copilot":
        api_key = os.getenv("GITHUB_TOKEN", "")
        base_url = "https://models.inference.ai.azure.com"
        model = LLM_VISION_MODEL
    elif provider in {"local_openai", "local", "openai-compatible", "openai_compatible", "qwen", "qwen-local"}:
        api_key = os.getenv("LOCAL_LLM_API_KEY", "dummy")
        base_url = os.getenv("LOCAL_LLM_BASE_URL", "").strip()
        model = os.getenv("LOCAL_LLM_MODEL", LLM_VISION_MODEL)
    else:
        raise RuntimeError(f"OCR_PROVIDER/LLM_PROVIDER non supporté pour la vision : {provider}")

    if not api_key:
        raise RuntimeError("Clé API manquante pour l'OCR Vision")
    if provider in {"local_openai", "local", "openai-compatible", "openai_compatible", "qwen", "qwen-local"} and not base_url:
        raise RuntimeError("LOCAL_LLM_BASE_URL est requis pour l'OCR Vision local.")

    kwargs = {"api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAI(**kwargs), model


def _guess_mime_type(file_path: str) -> str:
    ext = os.path.splitext(file_path.lower())[1]
    if ext in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if ext == ".png":
        return "image/png"
    if ext == ".webp":
        return "image/webp"
    return "image/png"


def _encode_file(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def _pdf_to_images(pdf_path: Path, temp_dir: str) -> list[str]:
    """Render PDF pages to PNG images for Vision OCR."""
    try:
        import fitz  # PyMuPDF
    except Exception as e:
        raise RuntimeError(
            "PyMuPDF requis pour le rendu PDF->images. "
            "Installez pymupdf puis relancez. "
            f"Détail: {e}"
        )

    doc = fitz.open(str(pdf_path))
    paths = []
    for i, page in enumerate(doc, start=1):
        mat = fitz.Matrix(2.5, 2.5)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        image_path = os.path.join(temp_dir, f"page_{i}.png")
        pix.save(image_path)
        paths.append(image_path)
    doc.close()
    return paths


def _extract_one_image_with_llm(client: OpenAI, model: str, image_path: str, page_index: int) -> str:
    mime_type = _guess_mime_type(image_path)
    base64_content = _encode_file(image_path)

    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Tu fais uniquement de l'OCR. Extrais tout le texte visible de cette page de CV. "
                            "Conserve au mieux l'ordre de lecture. Aucune explication."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{base64_content}"},
                    },
                ],
            }
        ],
        temperature=0,
        max_tokens=4096,
    )

    content = response.choices[0].message.content
    text = content if isinstance(content, str) else str(content)
    return f"\n===== PAGE {page_index} =====\n{text.strip()}\n"


def _ocr_pdf_via_api(file_path: Path) -> str:
    """OCR d'un PDF via API Vision (rendu pages -> image -> LLM)."""
    client, model = _get_vision_client()

    with tempfile.TemporaryDirectory(prefix="ocr_llm_pages_") as temp_dir:
        image_paths = _pdf_to_images(file_path, temp_dir)
        if not image_paths:
            raise RuntimeError("Aucune page extraite du PDF")

        chunks = []
        for idx, img_path in enumerate(image_paths, start=1):
            logger.info("OCR Vision page %d/%d", idx, len(image_paths))
            chunks.append(_extract_one_image_with_llm(client, model, img_path, idx))
        return "\n".join(chunks).strip()


def _ocr_pdf(file_path: Path) -> str:
    """OCR d'un PDF scanné page par page via pytesseract (utilise PyMuPDF pour le rendu)."""
    try:
        import pytesseract
        from PIL import Image
        import fitz
        from outils.image_processor import preprocess_for_ocr

        # Configuration Tesseract pour Windows si non présent dans le PATH
        if os.name == 'nt':
            tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_cmd):
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            else:
                logger.warning("Tesseract non trouvé à l'emplacement par défaut sur Windows.")
    except ImportError:
        raise RuntimeError(
            "PyMuPDF, pytesseract et opencv sont requis pour l'OCR. "
            "Installe-les avec : pip install pymupdf pytesseract opencv-python\n"
        )

    doc = fitz.open(str(file_path))
    pages_text = []

    # Dossier de debug pour les images binarisées
    debug_dir = Path("debug_ocr")
    debug_dir.mkdir(exist_ok=True)

    for page_num, page in enumerate(doc):
        # Rendu en haute résolution (300 DPI)
        pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72)) 
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # --- PHASE 3 : NETTOYAGE ADAPTATIF ---
        # Nettoyage de l'image (Grayscale, CLAHE, Adaptive Threshold) avant OCR
        img = preprocess_for_ocr(img)

        # Sauvegarde de l'image binarisée pour debug
        debug_path = debug_dir / f"{file_path.stem}_p{page_num+1}_bin.png"
        img.save(debug_path)
        
        # --- PHASE 4 : SEGMENTATION & OCR ---
        # Au lieu de lire toute la page, on détecte les blocs pour respecter l'ordre (colonnes)
        from outils.image_processor import get_text_blocks
        blocks = get_text_blocks(img)
        
        if not blocks:
            # Fallback si aucun bloc détecté
            text = pytesseract.image_to_string(img, lang="fra+eng", config="--psm 1")
            pages_text.append(text)
        else:
            page_content = []
            img_np = np.array(img)
            for (x, y, w, h) in blocks:
                # Extraire la zone du bloc
                roi = img_np[y:y+h, x:x+w]
                # OCR sur le bloc uniquement (PSM 6 : bloc de texte uniforme)
                block_text = pytesseract.image_to_string(roi, lang="fra+eng", config="--psm 6")
                page_content.append(block_text.strip())
            
            pages_text.append("\n\n".join(page_content))
            logger.info("OCR page %d : %d blocs détectés", page_num + 1, len(blocks))
    
    doc.close()
    return "\n".join(pages_text)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parse_docx(file_path: Path) -> str:
    """Extrait le texte d'un fichier DOCX (paragraphes + tableaux)."""
    doc = Document(str(file_path))
    parts = []

    # Paragraphes
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tableaux (souvent utilisés dans les CVs pour les compétences / dates)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Nettoyage texte
# ---------------------------------------------------------------------------

def _clean_text(text: str, is_markdown: bool = False) -> str:
    """Nettoie le texte extrait : artefacts, espaces multiples, lignes vides."""
    if not text:
        return ""

    # Remplacer les caractères de contrôle (sauf newline et tab)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Normaliser les fins de ligne Windows
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    if not is_markdown:
        # Supprimer les espaces en fin de ligne
        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
        # Supprimer les espaces multiples (garder un seul)
        text = re.sub(r"[ \t]{2,}", " ", text)
        # Recoudre les mots coupés en fin de ligne (tiret + saut de ligne)
        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    else:
        # En Markdown, on est plus conservateur sur les espaces pour ne pas briser les tableaux
        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)

    # Réduire les lignes vides multiples (max 2 consécutives)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


_QUALITY_HINT_PATTERNS = (
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    r"(?:\+?\d[\d\s()./\-]{7,}\d)",
    r"\b(?:experience|experiences|formation|education|competences|skills|profil|profile|summary)\b",
)


def _compute_text_quality(text: str, page_count: int | None = None) -> dict:
    """Calcule des métriques simples de qualité du texte extrait."""
    stripped = text or ""
    char_count = len(stripped)
    non_ws_chars = sum(1 for c in stripped if not c.isspace())
    alpha_count = sum(1 for c in stripped if c.isalpha())
    alnum_count = sum(1 for c in stripped if c.isalnum())
    alpha_ratio = alpha_count / max(alnum_count, 1)
    word_count = len(re.findall(r"\S+", stripped))
    line_count = stripped.count("\n") + 1 if stripped else 0
    effective_pages = max(page_count or line_count or 1, 1)
    chars_per_page = char_count / effective_pages
    words_per_page = word_count / effective_pages
    normalized = stripped.lower()
    anchor_hits = sum(1 for pattern in _QUALITY_HINT_PATTERNS if re.search(pattern, normalized, flags=re.IGNORECASE))

    return {
        "char_count": char_count,
        "non_whitespace_count": non_ws_chars,
        "word_count": word_count,
        "line_count": line_count,
        "page_count": effective_pages,
        "chars_per_page": round(chars_per_page, 2),
        "words_per_page": round(words_per_page, 2),
        "anchor_hits": anchor_hits,
        "alpha_ratio": round(alpha_ratio, 4),
    }


def _is_quality_sufficient(quality: dict) -> bool:
    """Détermine si la qualité du texte est suffisante pour le pipeline."""
    legacy_only = not any(
        key in quality for key in ("word_count", "words_per_page", "chars_per_page", "anchor_hits")
    )
    if legacy_only:
        return (
            quality.get("char_count", 0) >= OCR_MIN_CHARS
            and quality.get("alpha_ratio", 0.0) >= OCR_MIN_ALPHA_RATIO
        )

    return (
        quality.get("char_count", 0) >= OCR_MIN_CHARS
        and quality.get("alpha_ratio", 0.0) >= OCR_MIN_ALPHA_RATIO
        and (
            quality.get("word_count", 0) >= OCR_MIN_WORDS
            or quality.get("words_per_page", 0.0) >= OCR_MIN_WORDS_PER_PAGE
            or quality.get("chars_per_page", 0.0) >= MIN_CHARS_PER_PAGE
            or quality.get("anchor_hits", 0) >= OCR_MIN_ANCHOR_HITS
        )
    )


def _quality_rank_tuple(quality: dict) -> tuple[bool, int, float, float, int, int]:
    return (
        _is_quality_sufficient(quality),
        int(quality.get("anchor_hits", 0)),
        int(round(float(quality.get("words_per_page", 0.0)))),
        float(quality.get("alpha_ratio", 0.0)),
        int(quality.get("word_count", 0)),
        int(quality.get("char_count", 0)),
    )


def _best_quality_candidate(candidates: list[tuple[str, str, dict]]) -> tuple[str, str, dict]:
    available = [candidate for candidate in candidates if candidate[0]]
    if not available:
        return "", "", _compute_text_quality("")
    return max(available, key=lambda item: _quality_rank_tuple(item[2]))


# ---------------------------------------------------------------------------
# Utilitaire CLI pour test rapide
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    import json

    logging.basicConfig(level=logging.DEBUG, format="%(levelname)s | %(message)s")

    if len(sys.argv) < 2:
        print("Usage: python -m skills.parser <fichier.pdf|fichier.docx>")
        sys.exit(1)

    result = parse_file(sys.argv[1])
    # Afficher les métadonnées (sans le texte complet pour la lisibilité)
    meta = {k: v for k, v in result.items() if k != "text"}
    print(json.dumps(meta, ensure_ascii=False, indent=2))
    print(f"\n--- Texte extrait (premiers 500 caractères) ---\n{result['text'][:500]}")
